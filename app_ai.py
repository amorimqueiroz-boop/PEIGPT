import streamlit as st
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt
from openai import OpenAI
from pypdf import PdfReader
from fpdf import FPDF
import base64
import json
import os
import re
import glob

# ==============================================================================
# 1. CONFIGURAÇÃO INICIAL
# ==============================================================================
def get_favicon():
    if os.path.exists("iconeaba.png"): return "iconeaba.png"
    return "📘"

st.set_page_config(
    page_title="PEI 360º",
    page_icon=get_favicon(),
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==============================================================================
# 2. DESIGN SYSTEM "AWARD WINNING" (CSS AVANÇADO)
# ==============================================================================
def aplicar_estilo_visual():
    estilo = """
    <style>
        /* IMPORTANDO FONTE PREMIUM (INTER) */
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700;800&display=swap');
        
        /* VARIÁVEIS GLOBAIS */
        :root {
            --primary: #004E92;
            --secondary: #FF6B6B;
            --bg-app: #F8F9FD;
            --bg-card: #FFFFFF;
            --text-main: #1A202C;
            --text-light: #718096;
            --radius: 12px;
            --shadow-soft: 0 4px 20px rgba(0, 0, 0, 0.03);
            --shadow-hover: 0 10px 25px rgba(0, 78, 146, 0.1);
        }

        /* RESET GERAL */
        html, body, [class*="css"] {
            font-family: 'Inter', sans-serif;
            color: var(--text-main);
            background-color: var(--bg-app);
        }

        /* FUNDO DO APP */
        .stApp {
            background-color: var(--bg-app);
        }

        /* --- CABEÇALHO UNIFICADO (HERO SECTION) --- */
        .header-unified {
            background: linear-gradient(135deg, #FFFFFF 0%, #F8F9FD 100%);
            padding: 30px 40px;
            border-radius: 16px;
            border: 1px solid rgba(255,255,255,0.8);
            box-shadow: var(--shadow-soft);
            margin-bottom: 30px;
            display: flex;
            align-items: center;
            gap: 25px;
            position: relative;
            overflow: hidden;
        }
        .header-unified::before {
            content: '';
            position: absolute;
            top: 0; left: 0; width: 6px; height: 100%;
            background: var(--primary);
        }
        .header-unified p {
            color: var(--primary);
            margin: 0;
            font-size: 1.8rem;
            font-weight: 800;
            letter-spacing: -0.5px;
        }

        /* --- ABAS (TABS) --- */
        .stTabs [data-baseweb="tab-list"] {
            gap: 8px;
            padding: 10px 0 20px 0;
            background-color: transparent;
        }
        .stTabs [data-baseweb="tab"] {
            height: 45px;
            border-radius: 8px;
            padding: 0 20px;
            background-color: white;
            border: 1px solid #E2E8F0;
            font-weight: 600;
            color: var(--text-light);
            font-size: 0.9rem;
            transition: all 0.3s ease;
        }
        .stTabs [aria-selected="true"] {
            background-color: var(--primary) !important;
            color: white !important;
            border-color: var(--primary) !important;
            box-shadow: 0 4px 12px rgba(0, 78, 146, 0.3);
        }

        /* --- INPUTS & CAMPOS (FORMULÁRIOS ELEGANTES) --- */
        .stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"], .stDateInput input {
            background-color: white !important;
            border: 1px solid #E2E8F0 !important;
            border-radius: 10px !important;
            padding: 10px 12px !important;
            font-size: 0.95rem;
            color: var(--text-main) !important;
            transition: border 0.2s ease;
        }
        .stTextInput input:focus, .stTextArea textarea:focus {
            border-color: var(--primary) !important;
            box-shadow: 0 0 0 2px rgba(0, 78, 146, 0.1);
        }

        /* --- BOTÕES --- */
        div[data-testid="column"] .stButton button {
            border-radius: 10px !important;
            font-weight: 700 !important;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            height: 48px !important;
            border: none;
            transition: transform 0.2s, box-shadow 0.2s;
        }
        div[data-testid="column"] .stButton button:hover {
            transform: translateY(-2px);
            box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        }

        /* --- CONTAINERS (CARDS) --- */
        [data-testid="stVerticalBlock"] > [style*="flex-direction: column;"] > [data-testid="stVerticalBlock"] {
            /* Isso afeta containers internos */
        }
        
        /* CUSTOM CARDS PARA HOME */
        .rich-card {
            background-color: white;
            padding: 25px;
            border-radius: 16px;
            border: 1px solid #EDF2F7;
            box-shadow: var(--shadow-soft);
            transition: all 0.3s ease;
            height: 200px;
            display: flex;
            flex-direction: column;
            justify-content: center;
            text-decoration: none;
            color: inherit;
            position: relative;
        }
        .rich-card:hover {
            transform: translateY(-5px);
            border-color: var(--primary);
            box-shadow: var(--shadow-hover);
        }
        .rich-card h3 {
            color: var(--primary);
            font-size: 1.1rem;
            font-weight: 800;
            margin-bottom: 8px;
        }
        .rich-card p {
            color: var(--text-light);
            font-size: 0.9rem;
            line-height: 1.5;
        }

        /* --- AJUSTES DE SLIDERS --- */
        div[data-baseweb="slider"] {
            padding-top: 10px;
        }

        /* --- TITULOS DE SEÇÃO --- */
        h3 {
            font-weight: 800 !important;
            letter-spacing: -0.5px;
            margin-bottom: 20px !important;
            color: #2D3748;
        }
        
        /* SEPARADORES */
        hr {
            margin: 30px 0;
            border-color: #EDF2F7;
        }
    </style>
    <link href="https://cdn.jsdelivr.net/npm/remixicon@4.1.0/fonts/remixicon.css" rel="stylesheet">
    """
    st.markdown(estilo, unsafe_allow_html=True)

aplicar_estilo_visual()

# ==============================================================================
# 3. LISTAS DE DADOS (PRESERVADAS)
# ==============================================================================
LISTA_SERIES = [
    "Educação Infantil", "1º Ano (Fund. I)", "2º Ano (Fund. I)", "3º Ano (Fund. I)", 
    "4º Ano (Fund. I)", "5º Ano (Fund. I)", "6º Ano (Fund. II)", "7º Ano (Fund. II)", 
    "8º Ano (Fund. II)", "9º Ano (Fund. II)", "1ª Série (Ensino Médio)", 
    "2ª Série (Ensino Médio)", "3ª Série (Ensino Médio)"
]

LISTAS_BARREIRAS = {
    "Cognitivo": ["Atenção Sustentada", "Atenção Alternada", "Memória de Trabalho", "Memória de Curto Prazo", "Controle Inibitório", "Flexibilidade Cognitiva", "Planejamento e Organização", "Velocidade de Processamento", "Raciocínio Lógico/Abstrato"],
    "Comunicacional": ["Linguagem Expressiva (Fala)", "Linguagem Receptiva (Compreensão)", "Vocabulário Restrito", "Pragmática (Uso Social)", "Articulação/Fonologia", "Comunicação Não-Verbal", "Necessidade de CAA"],
    "Socioemocional": ["Regulação Emocional", "Tolerância à Frustração", "Interação com Pares", "Interação com Adultos", "Compreensão de Regras Sociais", "Rigidez de Pensamento", "Autoestima", "Agressividade"],
    "Sensorial/Motor": ["Coordenação Motora Fina", "Coordenação Motora Ampla", "Hipersensibilidade Auditiva", "Hipersensibilidade Tátil", "Hipersensibilidade Visual", "Busca Sensorial", "Tônus Muscular", "Planejamento Motor"],
    "Acadêmico": ["Alfabetização (Decodificação)", "Compreensão Leitora", "Grafia/Legibilidade", "Produção Textual", "Raciocínio Lógico-Matemático", "Cálculo/Operações", "Resolução de Problemas"]
}

LISTA_POTENCIAS = ["Memória Visual", "Memória Auditiva", "Raciocínio Lógico", "Criatividade", "Habilidades Artísticas", "Musicalidade", "Tecnologia", "Hiperfoco", "Vocabulário Rico", "Empatia", "Liderança", "Esportes", "Persistência", "Curiosidade"]

LISTA_PROFISSIONAIS = ["Psicólogo", "Fonoaudiólogo", "Terapeuta Ocupacional", "Neuropediatra", "Psiquiatra", "Psicopedagogo", "Professor de Apoio", "AT"]

# ==============================================================================
# 4. GERENCIAMENTO DE ESTADO
# ==============================================================================
default_state = {
    'nome': '', 'nasc': date(2015, 1, 1), 'serie': None, 'turma': '', 'diagnostico': '', 
    'lista_medicamentos': [], 'composicao_familiar': '', 'historico': '', 'familia': '', 
    'hiperfoco': '', 'potencias': [], 'rede_apoio': [], 'orientacoes_especialistas': '',
    'checklist_evidencias': {}, 
    'barreiras_selecionadas': {k: [] for k in LISTAS_BARREIRAS.keys()},
    'niveis_suporte': {}, 
    'estrategias_acesso': [], 'estrategias_ensino': [], 'estrategias_avaliacao': [], 
    'ia_sugestao': '', 'outros_acesso': '', 'outros_ensino': '', 
    'monitoramento_data': None, 
    'status_meta': 'Não Iniciado', 'parecer_geral': 'Manter Estratégias', 'proximos_passos_select': []
}

if 'dados' not in st.session_state: st.session_state.dados = default_state
else:
    for key, val in default_state.items():
        if key not in st.session_state.dados: st.session_state.dados[key] = val

if 'pdf_text' not in st.session_state: st.session_state.pdf_text = ""

# ==============================================================================
# 5. UTILITÁRIOS E BANCO
# ==============================================================================
PASTA_BANCO = "banco_alunos"
if not os.path.exists(PASTA_BANCO): os.makedirs(PASTA_BANCO)

def finding_logo():
    possiveis = ["360.png", "360.jpg", "logo.png", "logo.jpg", "iconeaba.png"]
    for nome in possiveis:
        if os.path.exists(nome): return nome
    return None

def get_base64_image(image_path):
    if not image_path: return ""
    with open(image_path, "rb") as img_file: return base64.b64encode(img_file.read()).decode()

def ler_pdf(arquivo):
    try:
        reader = PdfReader(arquivo); texto = ""
        for i, page in enumerate(reader.pages):
            if i >= 6: break 
            texto += page.extract_text() + "\n"
        return texto
    except: return ""

def limpar_texto_pdf(texto):
    if not texto: return ""
    texto = texto.replace('**', '').replace('__', '').replace('### ', '').replace('## ', '').replace('# ', '')
    return re.sub(r'[^\x00-\xff]', '', texto)

def salvar_aluno(dados):
    if not dados['nome']: return False, "Nome obrigatório."
    nome_arq = re.sub(r'[^a-zA-Z0-9]', '_', dados['nome'].lower()) + ".json"
    try:
        with open(os.path.join(PASTA_BANCO, nome_arq), 'w', encoding='utf-8') as f:
            json.dump(dados, f, default=str, ensure_ascii=False, indent=4)
        return True, f"Salvo: {dados['nome']}"
    except Exception as e: return False, str(e)

def carregar_aluno(nome_arq):
    try:
        with open(os.path.join(PASTA_BANCO, nome_arq), 'r', encoding='utf-8') as f: d = json.load(f)
        if 'nasc' in d: d['nasc'] = date.fromisoformat(d['nasc'])
        if d.get('monitoramento_data'): d['monitoramento_data'] = date.fromisoformat(d['monitoramento_data'])
        return d
    except: return None

# ==============================================================================
# 6. INTELIGÊNCIA ARTIFICIAL
# ==============================================================================
def consultar_gpt_pedagogico(api_key, dados, contexto_pdf=""):
    if not api_key: return None, "⚠️ Configure a Chave API OpenAI na lateral."
    try:
        client = OpenAI(api_key=api_key)
        evid = "\n".join([f"- {k}" for k, v in dados['checklist_evidencias'].items() if v])
        map_txt = ""
        for c, i in dados['barreiras_selecionadas'].items():
            if i: map_txt += f"\n[{c}]: " + ", ".join([f"{x} ({dados['niveis_suporte'].get(f'{c}_{x}','Monitorado')})" for x in i])
        
        sys = """Você é Especialista em BNCC e Educação Inclusiva. 
        Gere um PEI técnico de 6 seções integrando:
        1. Habilidades do Ano Corrente (BNCC).
        2. Habilidades de Anos Anteriores para Recomposição.
        3. Cruzamento do Hiperfoco como facilitador de aprendizagem.
        Use CAIXA ALTA apenas nos títulos numerados."""
        
        usr = f"ALUNO: {dados['nome']} | SÉRIE: {dados['serie']}\nDIAG: {dados['diagnostico']}\nHIST: {dados['historico']}\nEVIDÊNCIAS: {evid}\nBARREIRAS: {map_txt}\nHIPERFOCO: {dados['hiperfoco']}\nESTRATÉGIAS: {dados['estrategias_ensino']}\nLAUDO: {contexto_pdf[:5000]}"
        
        res = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "system", "content": sys}, {"role": "user", "content": usr}])
        return res.choices[0].message.content, None
    except Exception as e: return None, str(e)

# ==============================================================================
# 7. GERADOR PDF
# ==============================================================================
class PDF_V3(FPDF):
    def header(self):
        self.set_draw_color(0, 78, 146); self.set_line_width(0.4); self.rect(5, 5, 200, 287)
        logo = finding_logo()
        if logo: self.image(logo, 10, 10, 30); x_offset = 45 
        else: x_offset = 12
        self.set_xy(x_offset, 16); self.set_font('Arial', 'B', 16); self.set_text_color(0, 78, 146)
        self.cell(0, 8, 'PLANO DE ENSINO INDIVIDUALIZADO', 0, 1, 'L')
        self.set_xy(x_offset, 23); self.set_font('Arial', 'I', 10); self.set_text_color(100); self.cell(0, 5, 'Documento Oficial', 0, 1, 'L'); self.ln(20)
    def footer(self):
        self.set_y(-15); self.set_font('Arial', 'I', 8); self.set_text_color(128); self.cell(0, 10, f'PEI 360º | Página {self.page_no()}', 0, 0, 'C')
    def section_title(self, label):
        self.ln(8); self.set_fill_color(240, 248, 255); self.set_text_color(0, 78, 146); self.set_font('Arial', 'B', 11); self.cell(0, 8, f"  {label}", 0, 1, 'L', fill=True); self.ln(4)

def gerar_pdf_final(dados):
    pdf = PDF_V3(); pdf.add_page(); pdf.set_auto_page_break(auto=True, margin=20)
    pdf.section_title("1. IDENTIFICAÇÃO")
    pdf.set_font("Arial", size=10); pdf.set_text_color(0)
    pdf.cell(40, 6, "Nome:", 0, 0); pdf.cell(0, 6, dados['nome'], 0, 1)
    pdf.cell(40, 6, "Série:", 0, 0); pdf.cell(0, 6, f"{dados['serie']} - {dados['turma']}", 0, 1)
    if dados['ia_sugestao']:
        for linha in dados['ia_sugestao'].split('\n'):
            if re.match(r'^[1-6]\.', linha.strip()) and linha.strip().isupper(): pdf.section_title(linha)
            else: pdf.multi_cell(0, 6, limpar_texto_pdf(linha))
    
    if dados.get('monitoramento_data'):
        pdf.section_title("CRONOGRAMA DE REVISÃO E MONITORAMENTO")
        pp = ', '.join(dados.get('proximos_passos_select', []))
        txt = f"Previsão de Revisão: {dados['monitoramento_data'].strftime('%d/%m/%Y')}\n\nStatus da Meta: {dados.get('status_meta','-')}\n\nParecer Geral: {dados.get('parecer_geral','-')}\n\nPróximos Passos: {pp}"
        pdf.multi_cell(0, 6, limpar_texto_pdf(txt))
    return pdf.output(dest='S').encode('latin-1', 'replace')

def gerar_docx_final(dados):
    doc = Document(); style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(11)
    doc.add_heading('PLANO DE ENSINO INDIVIDUALIZADO', 0)
    doc.add_paragraph(f"Estudante: {dados['nome']}")
    if dados['ia_sugestao']: doc.add_paragraph(dados['ia_sugestao'])
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0); return buffer

# ==============================================================================
# 8. INTERFACE UI (DESIGN PREMIUM APLICADO)
# ==============================================================================
# SIDEBAR
with st.sidebar:
    logo = finding_logo()
    if logo: st.image(logo, width=130) # Aumentei um pouco a logo
    
    api_key = st.text_input("Chave OpenAI:", type="password") if 'OPENAI_API_KEY' not in st.secrets else st.secrets['OPENAI_API_KEY']
    st.markdown("---")
    st.caption("📂 Gestão de Casos")
    st.info("Para salvar, use as opções de Rascunho na aba 'Documento'.")
    st.markdown("---")
    data_atual = date.today().strftime("%d/%m/%Y")
    st.markdown(f"<div style='font-size:0.75rem; color:#A0AEC0; text-align:center;'><b>PEI 360º v14.0</b><br>Design System Premium<br>Dev: <b>Rodrigo A. Queiroz</b><br>{data_atual}</div>", unsafe_allow_html=True)

# HEADER
logo_path = finding_logo(); b64_logo = get_base64_image(logo_path)
img_html = f'<img src="data:image/png;base64,{b64_logo}" style="height: 70px;">' if logo_path else ""
st.markdown(f'<div class="header-unified">{img_html}<div><p>Ecossistema de Inteligência Pedagógica e Inclusiva</p></div></div>', unsafe_allow_html=True)

# ABAS
abas = ["Início", "Estudante", "Coleta de Evidências", "Rede de Apoio", "Potencialidades & Barreiras", "Plano de Ação", "Monitoramento", "Consultoria IA", "Documento"]
tab0, tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs(abas)

with tab0: # INÍCIO
    if api_key:
        with st.spinner("Conectando à IA..."):
            saudacao = "Bem-vindo, Especialista!" # Fallback rápido
            noticia = "Dica: Personalize as práticas de ensino para maior engajamento."
    
        st.markdown(f"""
        <div style="background: linear-gradient(120deg, #004E92 0%, #000428 100%); padding: 25px; border-radius: 16px; color: white; margin-bottom: 25px; box-shadow: 0 10px 25px rgba(0,78,146,0.25);">
            <div style="display:flex; gap:20px; align-items:center;">
                <i class="ri-sparkling-fill" style="font-size: 2.2rem; color: #FFD700;"></i>
                <div><h3 style="color:white; margin:0; font-size: 1.4rem;">{saudacao}</h3><p style="margin:5px 0 0 0; opacity:0.9;">Pronto para transformar a inclusão hoje?</p></div>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("### <i class='ri-apps-2-line'></i> Fundamentos", unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns(4)
    with c1: st.markdown('<div class="rich-card"><h3><i class="ri-book-open-line" style="color:#004E92"></i> O que é PEI?</h3><p>Conceitos fundamentais da Educação Inclusiva.</p></div>', unsafe_allow_html=True)
    with c2: st.markdown('<div class="rich-card"><h3><i class="ri-scales-3-line" style="color:#004E92"></i> Legislação</h3><p>Lei Brasileira de Inclusão e Decretos.</p></div>', unsafe_allow_html=True)
    with c3: st.markdown('<div class="rich-card"><h3><i class="ri-brain-line" style="color:#004E92"></i> Neurociência</h3><p>Artigos sobre desenvolvimento atípico.</p></div>', unsafe_allow_html=True)
    with c4: st.markdown('<div class="rich-card"><h3><i class="ri-compass-3-line" style="color:#004E92"></i> BNCC</h3><p>Base Nacional Comum Curricular.</p></div>', unsafe_allow_html=True)

with tab1: # ESTUDANTE
    st.markdown("### <i class='ri-user-star-line'></i> Dossiê do Estudante", unsafe_allow_html=True)
    with st.container(border=True): # Card
        c1, c2, c3, c4 = st.columns([3, 2, 2, 1])
        st.session_state.dados['nome'] = c1.text_input("Nome Completo", st.session_state.dados['nome'])
        st.session_state.dados['nasc'] = c2.date_input("Nascimento", value=st.session_state.dados['nasc'])
        st.session_state.dados['serie'] = c3.selectbox("Série/Ano", LISTA_SERIES, placeholder="Selecione...")
        st.session_state.dados['turma'] = c4.text_input("Turma", st.session_state.dados['turma'])
        
        st.divider()
        c1, c2 = st.columns(2)
        st.session_state.dados['historico'] = c1.text_area("Histórico Escolar", st.session_state.dados['historico'], placeholder="Trajetória, retenções e avanços.")
        st.session_state.dados['familia'] = c2.text_area("Contexto Familiar", st.session_state.dados['familia'], placeholder="Dinâmica familiar e acompanhamento.")
        st.session_state.dados['diagnostico'] = st.text_input("Diagnóstico Clínico", st.session_state.dados['diagnostico'])

    st.write("")
    with st.container(border=True): # Card Medicação
        st.markdown("##### <i class='ri-medicine-bottle-line'></i> Controle de Medicação", unsafe_allow_html=True)
        c1, c2, c3 = st.columns([3, 2, 1])
        nm = c1.text_input("Nome Med", key="nm_med")
        pos = c2.text_input("Posologia", key="pos_med")
        if c3.button("➕ Adicionar"):
            st.session_state.dados['lista_medicamentos'].append({"nome": nm, "posologia": pos, "escola": False}); st.rerun()
        
        if st.session_state.dados['lista_medicamentos']:
            st.markdown("---")
            for i, m in enumerate(st.session_state.dados['lista_medicamentos']):
                c_a, c_b, c_c = st.columns([4, 2, 1])
                c_a.markdown(f"**{m['nome']}**")
                m['escola'] = c_b.checkbox("Tomar na Escola?", m['escola'], key=f"esc_{i}")
                if c_c.button("🗑️", key=f"del_{i}"): st.session_state.dados['lista_medicamentos'].pop(i); st.rerun()

with tab2: # EVIDÊNCIAS
    st.markdown("### <i class='ri-search-eye-line'></i> Coleta de Evidências", unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    # Cards verticais para cada categoria
    with c1:
        with st.container(border=True):
            st.markdown("##### Aprendizagem")
            for q in ["Estagnação na aprendizagem", "Dificuldade de generalização", "Dificuldade de abstração", "Lacuna em pré-requisitos"]:
                st.session_state.dados['checklist_evidencias'][q] = st.checkbox(q, st.session_state.dados['checklist_evidencias'].get(q, False))
    with c2:
        with st.container(border=True):
            st.markdown("##### Atenção & Foco")
            for q in ["Oscilação de foco", "Fadiga mental rápida", "Dificuldade de iniciar tarefas", "Esquecimento recorrente"]:
                st.session_state.dados['checklist_evidencias'][q] = st.checkbox(q, st.session_state.dados['checklist_evidencias'].get(q, False))
    with c3:
        with st.container(border=True):
            st.markdown("##### Comportamento")
            for q in ["Dependência de mediação (1:1)", "Baixa tolerância à frustração", "Desorganização de materiais", "Recusa de tarefas"]:
                st.session_state.dados['checklist_evidencias'][q] = st.checkbox(q, st.session_state.dados['checklist_evidencias'].get(q, False))

with tab3:
    st.markdown("### <i class='ri-team-line'></i> Rede de Apoio", unsafe_allow_html=True)
    with st.container(border=True):
        st.session_state.dados['rede_apoio'] = st.multiselect("Profissionais", LISTA_PROFISSIONAIS, st.session_state.dados['rede_apoio'], placeholder="Selecione...")
        st.session_state.dados['orientacoes_especialistas'] = st.text_area("Orientações Clínicas", st.session_state.dados['orientacoes_especialistas'], height=150)

with tab4: # MAPEAMENTO (LAYOUT FIXO + DESIGN PREMIUM)
    st.markdown("### <i class='ri-map-pin-user-line'></i> Mapeamento Integral", unsafe_allow_html=True)
    
    # CARD POTENCIALIDADES
    with st.container(border=True):
        st.markdown("#### <i class='ri-lightbulb-flash-line' style='color:#004E92'></i> Potencialidades e Hiperfoco", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        st.session_state.dados['hiperfoco'] = c1.text_input("Hiperfoco", st.session_state.dados['hiperfoco'], placeholder="Ex: Minecraft, Dinossauros...")
        p_val = [p for p in st.session_state.dados['potencias'] if p in LISTA_POTENCIAS]
        st.session_state.dados['potencias'] = c2.multiselect("Pontos Fortes", LISTA_POTENCIAS, default=p_val, placeholder="Selecione...")
    
    st.write("") # Espaço visual
    
    # CARD BARREIRAS (Layout Fixo 3 Colunas)
    with st.container(border=True):
        st.markdown("#### <i class='ri-barricade-line' style='color:#FF6B6B'></i> Barreiras e Nível de Suporte", unsafe_allow_html=True)
        
        c_bg1, c_bg2, c_bg3 = st.columns(3)
        
        def render_cat_premium(coluna, titulo, chave):
            with coluna:
                st.markdown(f"**{titulo}**")
                # Filtro de segurança
                opcoes = LISTAS_BARREIRAS[chave]
                salvos = [x for x in st.session_state.dados['barreiras_selecionadas'].get(chave, []) if x in opcoes]
                
                sel = st.multiselect("Selecione:", opcoes, key=f"ms_{chave}", default=salvos, placeholder="Selecione...", label_visibility="collapsed")
                st.session_state.dados['barreiras_selecionadas'][chave] = sel
                
                if sel:
                    for item in sel:
                        k = f"{chave}_{item}"
                        val = st.session_state.dados['niveis_suporte'].get(k, "Monitorado")
                        st.session_state.dados['niveis_suporte'][k] = st.select_slider(item, ["Autônomo", "Monitorado", "Substancial", "Muito Substancial"], value=val, key=f"sl_{k}")
                st.write("")

        # Coluna 1
        render_cat_premium(c_bg1, "Cognitivo", "Cognitivo")
        render_cat_premium(c_bg1, "Sensorial/Motor", "Sensorial/Motor")
        # Coluna 2
        render_cat_premium(c_bg2, "Comunicacional", "Comunicacional")
        render_cat_premium(c_bg2, "Acadêmico", "Acadêmico")
        # Coluna 3
        render_cat_premium(c_bg3, "Socioemocional", "Socioemocional")

with tab5:
    st.markdown("### <i class='ri-tools-line'></i> Plano de Ação Estratégico", unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    
    # Cards Individuais para cada eixo
    with c1:
        with st.container(border=True):
            st.markdown("##### 1. Acesso (DUA)")
            st.session_state.dados['estrategias_acesso'] = st.multiselect("Recursos:", ["Tempo Estendido", "Apoio Leitura"], default=st.session_state.dados['estrategias_acesso'], placeholder="Selecione...", key="acc")
            st.session_state.dados['outros_acesso'] = st.text_input("Prática Personalizada", st.session_state.dados['outros_acesso'], placeholder="Descreva aqui...")
    with c2:
        with st.container(border=True):
            st.markdown("##### 2. Ensino")
            st.session_state.dados['estrategias_ensino'] = st.multiselect("Metodologia:", ["Pistas Visuais", "Mapas Mentais"], default=st.session_state.dados['estrategias_ensino'], placeholder="Selecione...", key="ens")
            st.session_state.dados['outros_ensino'] = st.text_input("Prática Pedagógica", st.session_state.dados['outros_ensino'], placeholder="Descreva aqui...")
    with c3:
        with st.container(border=True):
            st.markdown("##### 3. Avaliação")
            st.session_state.dados['estrategias_avaliacao'] = st.multiselect("Formato:", ["Prova Adaptada", "Oral"], default=st.session_state.dados['estrategias_avaliacao'], placeholder="Selecione...", key="ava")

with tab6: # MONITORAMENTO
    st.markdown("### <i class='ri-loop-right-line'></i> Monitoramento", unsafe_allow_html=True)
    with st.container(border=True):
        st.info("Ciclo de revisão pedagógica do PEI.")
        c1, c2 = st.columns(2)
        st.session_state.dados['monitoramento_data'] = c1.date_input("Próxima Revisão", value=st.session_state.dados.get('monitoramento_data', date.today()))
        st.session_state.dados['status_meta'] = c2.selectbox("Status Atual", ["Não Iniciado", "Em Andamento", "Atingido"], placeholder="Selecione...")
        
        st.divider()
        c3, c4 = st.columns(2)
        st.session_state.dados['parecer_geral'] = c3.selectbox("Parecer Geral", ["Manter Estratégias", "Reduzir Suporte"], placeholder="Selecione...")
        st.session_state.dados['proximos_passos_select'] = c4.multiselect("Ações Futuras", ["Reunião Família", "Novo PEI"], placeholder="Selecione...")

with tab7:
    st.markdown("### <i class='ri-robot-2-line'></i> Consultoria IA", unsafe_allow_html=True)
    with st.container(border=True):
        if st.button("GERAR PLANO AGORA", type="primary"):
            res, err = consultar_gpt_pedagogico(api_key, st.session_state.dados, st.session_state.pdf_text)
            if res: st.session_state.dados['ia_sugestao'] = res; st.success("Plano Gerado!")
            else: st.error(err)
        
        if st.session_state.dados['ia_sugestao']:
            st.text_area("Relatório Editável", st.session_state.dados['ia_sugestao'], height=600)

with tab8:
    st.markdown("### <i class='ri-file-pdf-line'></i> Documento & Gestão", unsafe_allow_html=True)
    if st.session_state.dados['ia_sugestao']:
        c1, c2 = st.columns(2)
        with c1: st.download_button("📥 Baixar PDF Pro", gerar_pdf_final(st.session_state.dados), f"PEI_{st.session_state.dados['nome']}.pdf", "application/pdf")
        with c2:
            st.download_button("💾 Salvar Rascunho (JSON)", json.dumps(st.session_state.dados, default=str), f"PEI_{st.session_state.dados['nome']}.json", "application/json")
            up = st.file_uploader("Carregar Rascunho", type="json")
            if up:
                d = json.load(up); 
                if 'nasc' in d: d['nasc'] = date.fromisoformat(d['nasc'])
                st.session_state.dados.update(d); st.rerun()
    
    st.divider()
    st.markdown("#### 🗂️ Banco Local")
    for arq in glob.glob(os.path.join(PASTA_BANCO, "*.json")):
        nome = os.path.basename(arq).replace(".json", "").replace("_", " ").title()
        c1, c2, c3 = st.columns([6, 2, 2])
        c1.markdown(f"**{nome}**")
        if c2.button("📂 Abrir", key=f"open_{arq}"): st.session_state.dados = json.load(open(arq)); st.rerun()
    
    if st.button("Salvar no Banco Local"): salvar_aluno(st.session_state.dados); st.rerun()
