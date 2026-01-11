import streamlit as st
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt
from openai import OpenAI
from pypdf import PdfReader
from fpdf import FPDF
import base64
import json
import os
import re
import glob
import random
import requests

# ==============================================================================
# 1. CONFIGURAÇÃO INICIAL
# ==============================================================================
def get_favicon():
    return "🗺️"

st.set_page_config(
    page_title="PEI 360º Student Map",
    page_icon=get_favicon(),
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==============================================================================
# 2. LISTAS DE DADOS
# ==============================================================================
LISTA_SERIES = ["Educação Infantil", "1º Ano (Fund. I)", "2º Ano (Fund. I)", "3º Ano (Fund. I)", "4º Ano (Fund. I)", "5º Ano (Fund. I)", "6º Ano (Fund. II)", "7º Ano (Fund. II)", "8º Ano (Fund. II)", "9º Ano (Fund. II)", "1ª Série (EM)", "2ª Série (EM)", "3ª Série (EM)"]

LISTAS_BARREIRAS = {
    "Cognitivo": ["Atenção Sustentada", "Memória de Trabalho", "Flexibilidade Cognitiva", "Raciocínio Lógico"],
    "Comunicacional": ["Linguagem Expressiva", "Compreensão", "Pragmática (Uso Social)", "Vocabulário"],
    "Socioemocional": ["Regulação Emocional", "Tolerância à Frustração", "Interação Social", "Autoestima"],
    "Sensorial/Motor": ["Coordenação Motora", "Hipersensibilidade", "Busca Sensorial", "Planejamento Motor"],
    "Acadêmico": ["Alfabetização", "Compreensão Leitora", "Cálculo", "Produção Textual"]
}

LISTA_POTENCIAS = ["Memória Visual", "Musicalidade", "Tecnologia", "Hiperfoco", "Liderança", "Esportes", "Desenho", "Cálculo Mental", "Oralidade", "Criatividade"]
LISTA_PROFISSIONAIS = ["Psicólogo", "Fonoaudiólogo", "Terapeuta Ocupacional", "Neuropediatra", "Psiquiatra", "Psicopedagogo", "Professor de Apoio", "AT"]
LISTA_FAMILIA = ["Mãe", "Pai", "Mãe (2ª)", "Pai (2º)", "Avó", "Avô", "Irmão(s)", "Tio(a)", "Padrasto", "Madrasta", "Tutor Legal", "Abrigo Institucional"]

# ==============================================================================
# 3. GERENCIAMENTO DE ESTADO
# ==============================================================================
default_state = {
    'nome': '', 'nasc': date(2015, 1, 1), 'serie': None, 'turma': '', 'diagnostico': '', 
    'lista_medicamentos': [], 'composicao_familiar_tags': [], 'historico': '', 'familia': '', 
    'hiperfoco': '', 'potencias': [], 'rede_apoio': [], 'orientacoes_especialistas': '',
    'checklist_evidencias': {}, 
    'barreiras_selecionadas': {k: [] for k in LISTAS_BARREIRAS.keys()},
    'niveis_suporte': {}, 
    'estrategias_acesso': [], 'estrategias_ensino': [], 'estrategias_avaliacao': [], 
    'ia_sugestao': '', 'outros_acesso': '', 'outros_ensino': '', 
    'monitoramento_data': date.today(), 
    'status_meta': 'Não Iniciado', 'parecer_geral': 'Manter Estratégias', 'proximos_passos_select': [],
    'dalle_image_url': ''
}

if 'dados' not in st.session_state: st.session_state.dados = default_state
else:
    for key, val in default_state.items():
        if key not in st.session_state.dados: st.session_state.dados[key] = val

if 'dalle_image_url' not in st.session_state: st.session_state.dalle_image_url = ""
if 'pdf_text' not in st.session_state: st.session_state.pdf_text = ""

# ==============================================================================
# 4. LÓGICA E UTILITÁRIOS
# ==============================================================================
PASTA_BANCO = "banco_alunos"
if not os.path.exists(PASTA_BANCO): os.makedirs(PASTA_BANCO)

def calcular_idade(data_nasc):
    if not data_nasc: return ""
    hoje = date.today()
    idade = hoje.year - data_nasc.year - ((hoje.month, hoje.day) < (data_nasc.month, data_nasc.day))
    return f"{idade} anos"

def get_hiperfoco_emoji(texto):
    if not texto: return "🚀"
    t = texto.lower()
    if "jogo" in t or "game" in t or "minecraft" in t or "roblox" in t: return "🎮"
    if "dino" in t: return "🦖"
    if "fute" in t or "bola" in t: return "⚽"
    if "desenho" in t or "arte" in t: return "🎨"
    if "músic" in t: return "🎵"
    if "anim" in t or "gato" in t or "cachorro" in t: return "🐾"
    if "carro" in t: return "🏎️"
    if "espaço" in t: return "🪐"
    return "🚀"

def calcular_complexidade_pei(dados):
    n_bar = sum(len(v) for v in dados['barreiras_selecionadas'].values())
    n_suporte_alto = sum(1 for v in dados['niveis_suporte'].values() if v in ["Substancial", "Muito Substancial"])
    recursos = 0
    if dados['rede_apoio']: recursos += 3
    if dados['lista_medicamentos']: recursos += 2
    saldo = (n_bar + n_suporte_alto) - recursos
    if saldo <= 2: return "FLUIDA", "#F0FFF4", "#276749"
    if saldo <= 7: return "ATENÇÃO", "#FFFFF0", "#D69E2E"
    return "CRÍTICA", "#FFF5F5", "#C53030"

# --- EXTRAÇÃO DE TAGS (REFORÇADA) ---
def extrair_tag_ia(texto, tag):
    if not texto: return ""
    # Tenta encontrar a tag exata
    padrao = fr'\[{tag}\](.*?)(\[FIM_{tag}\]|\[|$)'
    match = re.search(padrao, texto, re.DOTALL | re.IGNORECASE)
    if match: 
        return match.group(1).strip()
    
    # Fallback para MAPA_TEXTO_GAMIFICADO se a tag falhar
    if tag == "MAPA_TEXTO_GAMIFICADO":
        if "⚡" in texto and "**" in texto:
             # Tenta pegar do primeiro emoji de raio até o fim ou próxima tag
             match_fallback = re.search(r'(⚡.*?)(?=\[|$)', texto, re.DOTALL)
             if match_fallback: return match_fallback.group(1).strip()
    
    return ""

def extrair_metas_estruturadas(texto):
    bloco = extrair_tag_ia(texto, "METAS_SMART")
    if not bloco: return None
    metas = {"Curto": "Definir...", "Medio": "Definir...", "Longo": "Definir..."}
    linhas = bloco.split('\n')
    for l in linhas:
        l_clean = re.sub(r'^[\-\*]+', '', l).strip()
        if "Curto" in l or "2 meses" in l: metas["Curto"] = l_clean.split(":")[-1].strip()
        elif "Médio" in l or "Semestre" in l: metas["Medio"] = l_clean.split(":")[-1].strip()
        elif "Longo" in l or "Ano" in l: metas["Longo"] = l_clean.split(":")[-1].strip()
    return metas

def extrair_bloom(texto):
    bloco = extrair_tag_ia(texto, "TAXONOMIA_BLOOM")
    if not bloco: return ["Identificar", "Compreender", "Aplicar"]
    return [v.strip() for v in bloco.split(',')]

def get_pro_icon(nome_profissional):
    p = nome_profissional.lower()
    if "psic" in p: return "🧠"
    if "fono" in p: return "🗣️"
    if "terapeuta" in p or "ocupacional" in p: return "🧩"
    if "neuro" in p: return "🩺"
    if "prof" in p: return "🍎"
    return "👨‍⚕️"

def finding_logo():
    possiveis = ["360.png", "360.jpg", "logo.png", "logo.jpg", "iconeaba.png"]
    for nome in possiveis:
        if os.path.exists(nome): return nome
    return None

def get_base64_image(image_path):
    if not image_path: return ""
    with open(image_path, "rb") as img_file: return base64.b64encode(img_file.read()).decode()

def ler_pdf(arquivo):
    try:
        reader = PdfReader(arquivo); texto = ""
        for i, page in enumerate(reader.pages):
            if i >= 6: break 
            texto += page.extract_text() + "\n"
        return texto
    except: return ""

def limpar_texto_pdf(texto):
    if not texto: return ""
    t = re.sub(r'\[.*?\]', '', texto) 
    t = t.replace('**', '').replace('__', '').replace('### ', '').replace('## ', '').replace('# ', '')
    return re.sub(r'[^\x00-\xff]', '', t)

def extrair_linhas_bncc(texto):
    padrao = r'([A-Z]{2}\d{1,2}[A-Z]{2,3}\d{2,3}.*?)(?=\n|$)'
    if not texto: return []
    linhas = re.findall(padrao, texto)
    return list(set([l.strip().replace('**', '') for l in linhas if len(l) > 10]))

def salvar_aluno(dados):
    if not dados['nome']: return False, "Nome obrigatório."
    nome_arq = re.sub(r'[^a-zA-Z0-9]', '_', dados['nome'].lower()) + ".json"
    try:
        with open(os.path.join(PASTA_BANCO, nome_arq), 'w', encoding='utf-8') as f:
            json.dump(dados, f, default=str, ensure_ascii=False, indent=4)
        return True, f"Salvo: {dados['nome']}"
    except Exception as e: return False, str(e)

def carregar_aluno(nome_arq):
    try:
        with open(os.path.join(PASTA_BANCO, nome_arq), 'r', encoding='utf-8') as f: d = json.load(f)
        if 'nasc' in d: d['nasc'] = date.fromisoformat(d['nasc'])
        if d.get('monitoramento_data'): d['monitoramento_data'] = date.fromisoformat(d['monitoramento_data'])
        return d
    except: return None

def excluir_aluno(nome_arq):
    try: os.remove(os.path.join(PASTA_BANCO, nome_arq)); return True
    except: return False

def calcular_progresso():
    if st.session_state.dados['ia_sugestao']: return 100
    return 50

def render_progresso():
    p = calcular_progresso()
    icon = "🌱"
    bar_color = "linear-gradient(90deg, #FF6B6B 0%, #FF8E53 100%)"
    if p >= 100: 
        icon = "🏆"
        bar_color = "linear-gradient(90deg, #00C6FF 0%, #0072FF 100%)" 
    st.markdown(f"""<div class="prog-container"><div class="prog-track"><div class="prog-fill" style="width: {p}%; background: {bar_color};"></div></div><div class="prog-icon" style="left: {p}%;">{icon}</div></div>""", unsafe_allow_html=True)

# ==============================================================================
# 5. ESTILO VISUAL
# ==============================================================================
def aplicar_estilo_visual():
    estilo = """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Nunito:wght@400;600;700;800&display=swap');
        html, body, [class*="css"] { font-family: 'Nunito', sans-serif; color: #2D3748; }
        .block-container { padding-top: 1rem !important; padding-bottom: 5rem !important; }
        div[data-baseweb="tab-border"], div[data-baseweb="tab-highlight"] { display: none !important; }
        
        .header-unified { background-color: white; padding: 20px 40px; border-radius: 16px; border: 1px solid #E2E8F0; box-shadow: 0 4px 15px rgba(0,0,0,0.03); margin-bottom: 20px; display: flex; align-items: center; gap: 20px; }
        .header-subtitle { color: #718096; font-size: 1.1rem; font-weight: 700; margin: 0; letter-spacing: 0.5px; border-left: 2px solid #E2E8F0; padding-left: 15px; }

        .stTabs [data-baseweb="tab-list"] { gap: 8px; flex-wrap: wrap; margin-bottom: 20px; justify-content: center; }
        .stTabs [data-baseweb="tab"] { height: 36px; border-radius: 18px !important; background-color: white; border: 1px solid #E2E8F0; color: #718096; font-weight: 700; font-size: 0.85rem; padding: 0 20px; transition: all 0.2s ease; }
        .stTabs [aria-selected="true"] { background-color: #FF6B6B !important; color: white !important; border-color: #FF6B6B !important; box-shadow: 0 4px 10px rgba(255, 107, 107, 0.3); }
        
        .prog-container { width: 100%; position: relative; margin: 0 0 40px 0; }
        .prog-track { width: 100%; height: 3px; background-color: #E2E8F0; border-radius: 1.5px; }
        .prog-fill { height: 100%; border-radius: 1.5px; transition: width 1.5s cubic-bezier(0.4, 0, 0.2, 1), background 1.5s ease; box-shadow: 0 1px 4px rgba(0,0,0,0.1); }
        .prog-icon { position: absolute; top: -23px; font-size: 1.8rem; transition: left 1.5s cubic-bezier(0.4, 0, 0.2, 1); transform: translateX(-50%); z-index: 10; filter: drop-shadow(0 2px 2px rgba(0,0,0,0.15)); }

        .dash-hero { background: linear-gradient(135deg, #0F52BA 0%, #062B61 100%); border-radius: 16px; padding: 25px; color: white; margin-bottom: 20px; display: flex; justify-content: space-between; align-items: center; box-shadow: 0 8px 15px rgba(15, 82, 186, 0.2); }
        .apple-avatar { width: 60px; height: 60px; border-radius: 50%; background: rgba(255,255,255,0.15); border: 2px solid rgba(255,255,255,0.4); color: white; font-weight: 800; font-size: 1.6rem; display: flex; align-items: center; justify-content: center; }

        .metric-card { background: white; border-radius: 16px; padding: 15px; border: 1px solid #E2E8F0; display: flex; flex-direction: column; align-items: center; justify-content: center; height: 140px; box-shadow: 0 2px 5px rgba(0,0,0,0.02); }
        .css-donut { width: 70px; height: 70px; border-radius: 50%; background: conic-gradient(var(--fill) var(--p), #EDF2F7 0); display: flex; align-items: center; justify-content: center; margin-bottom: 8px; }
        .css-donut::after { content: ""; position: absolute; width: 54px; height: 54px; border-radius: 50%; background: white; }
        .d-val { position: absolute; z-index: 2; font-size: 1.3rem; font-weight: 800; color: #2D3748; }
        .d-lbl { text-transform: uppercase; font-size: 0.65rem; color: #718096; font-weight: 700; letter-spacing: 0.5px; text-align: center; }
        .comp-icon-box { margin-bottom: 5px; }

        .soft-card { border-radius: 12px; padding: 20px; min-height: 220px; height: 100%; display: flex; flex-direction: column; box-shadow: 0 2px 5px rgba(0,0,0,0.02); border: 1px solid rgba(0,0,0,0.05); border-left: 5px solid; position: relative; overflow: hidden; }
        .sc-orange { background-color: #FFF5F5; border-left-color: #DD6B20; }
        .sc-blue { background-color: #EBF8FF; border-left-color: #3182CE; }
        .sc-yellow { background-color: #FFFFF0; border-left-color: #D69E2E; }
        .sc-cyan { background-color: #E6FFFA; border-left-color: #0BC5EA; }
        .sc-green { background-color: #F0FFF4; border-left-color: #38A169; }
        .sc-head { font-size: 0.75rem; font-weight: 800; text-transform: uppercase; margin-bottom: 12px; display: flex; align-items: center; gap: 8px; color: #4A5568; letter-spacing: 0.5px; z-index: 2; }
        .sc-body { font-size: 0.9rem; line-height: 1.6; color: #2D3748; font-weight: 600; z-index: 2; flex-grow: 1; }
        .bg-icon { position: absolute; bottom: -10px; right: -10px; font-size: 6rem; opacity: 0.08; z-index: 1; pointer-events: none; }
        
        .home-card { background-color: white; padding: 30px 20px; border-radius: 16px; border: 1px solid #E2E8F0; box-shadow: 0 4px 6px rgba(0,0,0,0.02); transition: all 0.3s ease; height: 250px; display: flex; flex-direction: column; align-items: center; justify-content: center; text-align: center; }
        .home-card:hover { transform: translateY(-5px); box-shadow: 0 15px 30px rgba(15, 82, 186, 0.1); border-color: #BEE3F8;}
        .home-card h3 { margin: 15px 0 10px 0; font-size: 1.1rem; color: #0F52BA; font-weight: 800; }
        .home-card p { font-size: 0.85rem; color: #718096; line-height: 1.4; margin: 0; }
        .icon-box { width: 70px; height: 70px; border-radius: 18px; display: flex; align-items: center; justify-content: center; font-size: 2.2rem; margin-bottom: 15px; }
        .ic-blue { background-color: #EBF8FF !important; color: #3182CE !important; border: 1px solid #BEE3F8 !important; }
        .ic-gold { background-color: #FFFFF0 !important; color: #D69E2E !important; border: 1px solid #FAF089 !important; }
        .ic-pink { background-color: #FFF5F7 !important; color: #D53F8C !important; border: 1px solid #FED7E2 !important; }
        .ic-green { background-color: #F0FFF4 !important; color: #38A169 !important; border: 1px solid #C6F6D5 !important; }
        .rich-card-link { text-decoration: none; color: inherit; display: block; height: 100%; }
        
        .rede-chip { display: inline-flex; align-items: center; background: white; padding: 6px 12px; border-radius: 20px; margin: 4px; box-shadow: 0 2px 4px rgba(0,0,0,0.05); font-size: 0.85rem; font-weight: 700; color: #2C5282; }
        .dna-bar-container { margin-bottom: 12px; }
        .dna-bar-flex { display: flex; justify-content: space-between; font-size: 0.8rem; margin-bottom: 4px; color: #4A5568; font-weight: 600; }
        .dna-bar-bg { width: 100%; height: 6px; background: #E2E8F0; border-radius: 3px; overflow: hidden; }
        .dna-bar-fill { height: 100%; border-radius: 3px; transition: width 0.5s ease; }
        .bloom-tag { background: #EBF8FF; color: #3182CE; padding: 4px 10px; border-radius: 12px; font-size: 0.8rem; font-weight: 700; margin-right: 5px; border: 1px solid #BEE3F8; display: inline-block; margin-bottom: 5px; }
        .meta-row { display: flex; align-items: center; gap: 10px; margin-bottom: 8px; font-size: 0.85rem; border-bottom: 1px solid rgba(0,0,0,0.05); padding-bottom: 5px; }
        
        .stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"], .stMultiSelect div[data-baseweb="select"] { border-radius: 10px !important; border-color: #E2E8F0 !important; }
        div[data-testid="column"] .stButton button { border-radius: 10px !important; font-weight: 800 !important; height: 50px !important; background-color: #0F52BA !important; color: white !important; border: none !important; }
        div[data-testid="column"] .stButton button:hover { background-color: #0A3D8F !important; }
        div[data-baseweb="checkbox"] div[class*="checked"] { background-color: #0F52BA !important; border-color: #0F52BA !important; }
        .ia-side-box { background: #F8FAFC; border-radius: 16px; padding: 25px; border: 1px solid #E2E8F0; text-align: left; margin-bottom: 20px; }
        .form-section-title { display: flex; align-items: center; gap: 10px; color: #0F52BA; font-weight: 700; font-size: 1.1rem; margin-top: 20px; margin-bottom: 15px; border-bottom: 2px solid #F7FAFC; padding-bottom: 5px; }
    </style>
    <link href="https://cdn.jsdelivr.net/npm/remixicon@4.1.0/fonts/remixicon.css" rel="stylesheet">
    """
    st.markdown(estilo, unsafe_allow_html=True)

aplicar_estilo_visual()

# ==============================================================================
# 6. INTELIGÊNCIA ARTIFICIAL (V90 - PROMPT REFORÇADO + REGEN)
# ==============================================================================
@st.cache_data(ttl=3600)
def gerar_saudacao_ia(api_key):
    if not api_key: return "Bem-vindo ao PEI 360º."
    try:
        client = OpenAI(api_key=api_key)
        res = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": "Frase curta inspiradora para professor sobre inclusão."}], temperature=0.9)
        return res.choices[0].message.content
    except: return "A inclusão transforma vidas."

@st.cache_data(ttl=3600)
def gerar_noticia_ia(api_key):
    if not api_key: return "Dica: Mantenha o PEI sempre atualizado."
    try:
        client = OpenAI(api_key=api_key)
        res = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": "Dica curta sobre legislação de inclusão ou neurociência (máx 2 frases)."}], temperature=0.7)
        return res.choices[0].message.content
    except: return "O cérebro aprende durante toda a vida."

# --- FUNÇÃO DALL-E 3 (INTEGRADA AO TEXTO) ---
def gerar_imagem_dalle_integrada(api_key, dados_aluno, texto_estrategias):
    if not api_key: return None, "Configure a API Key."
    if not texto_estrategias: return None, "Texto das estratégias não encontrado."
    try:
        client = OpenAI(api_key=api_key)
        hf = dados_aluno['hiperfoco'] if dados_aluno['hiperfoco'] else "aprendizado criativo"
        
        prompt_dalle = f"""
        A creative, colorful infographic illustration of a 'Student Power Map' pinned on a corkboard.
        Title area: "MEU MAPA DE PODER".
        Theme: {hf} (use visual elements from this theme).
        Content: Visual representations of these specific strategies:
        ---
        {texto_estrategias[:600]}
        ---
        Style: Pixar/Disney animation style, vibrant, organized, friendly.
        Layout: Central core with 4-5 colorful branches connecting to sticky notes or drawings.
        Atmosphere: Empowering, clear, fun.
        """

        with st.spinner("🎨 A IA está desenhando o mapa com suas estratégias... (15s)"):
            response = client.images.generate(
                model="dall-e-3", prompt=prompt_dalle, size="1024x1024", quality="standard", n=1,
            )
        return response.data[0].url, None
    except Exception as e:
        return None, str(e)

def consultar_gpt_pedagogico(api_key, dados, contexto_pdf="", regenerar=False):
    if not api_key: return None, "⚠️ Configure a Chave API."
    try:
        client = OpenAI(api_key=api_key)
        familia = ", ".join(dados['composicao_familiar_tags']) if dados['composicao_familiar_tags'] else "Não informado"
        evid = "\n".join([f"- {k.replace('?', '')}" for k, v in dados['checklist_evidencias'].items() if v])
        meds_info = "Nenhuma medicação informada."
        if dados['lista_medicamentos']:
            meds_info = "\n".join([f"- {m['nome']} ({m['posologia']}). Admin Escola: {'Sim' if m.get('escola') else 'Não'}." for m in dados['lista_medicamentos']])

        # Se for regeneração, pede algo diferente
        extra_instruction = ""
        if regenerar:
            extra_instruction = " (ATENÇÃO: O usuário pediu para REGENERAR. Use uma abordagem diferente, novas estratégias e um tom mais motivador)."

        prompt_sys = f"""
        Você é um Especialista Sênior em Neuroeducação, Inclusão e Legislação.{extra_instruction}
        SUA MISSÃO: Criar um PEI Técnico e um CONTEÚDO GAMIFICADO EM TEXTO PARA O ALUNO.
        
        --- TAGS OBRIGATÓRIAS ---
        [ANALISE_FARMA] ... [FIM_ANALISE_FARMA]
        [TAXONOMIA_BLOOM] 3 verbos cognitivos. Ex: Identificar, Classificar [FIM_TAXONOMIA_BLOOM]
        [METAS_SMART] 
        - CURTO PRAZO (2 meses): ...
        - MÉDIO PRAZO (Semestre): ...
        - LONGO PRAZO (Ano): ...
        [FIM_METAS_SMART]
        [ESTRATEGIA_MASTER] ... [FIM_ESTRATEGIA_MASTER]
        [MATRIZ_BNCC] ... [FIM_MATRIZ_BNCC]
        
        [MAPA_TEXTO_GAMIFICADO]
        Gere um guia visual em Markdown para o aluno, exatamente neste formato:
        
        ⚡ **Meu Mapa de Poderes** ⚡
        
        🧠 **Super Foco (Aprendizado)**
        [Descreva como usar o hiperfoco para aprender]
        
        🌬️ **Calma Interior (Ansiedade)**
        [Dica prática para se acalmar]
        
        🕒 **Botão de Pausa (Em Sala)**
        [Estratégia de autorregulação]
        
        📁 **Mestre da Organização**
        [Dica de rotina]
        
        🚶‍♂️ **Recarga de Energia**
        [Dica de pausa]
        
        🤝 **Meus Aliados**
        [Liste família/professores como suporte]

        🎨 **Dica:** Você pode desenhar este mapa e colar adesivos!
        [FIM_MAPA_TEXTO_GAMIFICADO]
        
        ESTRUTURA GERAL:
        1. 🌟 AVALIAÇÃO DE REPERTÓRIO: Foco na potência.
        2. 🧩 DIRETRIZES DE ADAPTAÇÃO: DUA e Avaliação.
        """
        
        prompt_user = f"""
        ALUNO: {dados['nome']} | SÉRIE: {dados['serie']}
        DIAGNÓSTICO: {dados['diagnostico']}
        MEDICAÇÃO: {meds_info}
        HIPERFOCO: {dados['hiperfoco']}
        BARREIRAS: {json.dumps(dados['barreiras_selecionadas'], ensure_ascii=False)}
        EVIDÊNCIAS: {evid}
        FAMILIA: {familia}
        LAUDO: {contexto_pdf[:3000] if contexto_pdf else "Nenhum."}
        """
        
        res = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "system", "content": prompt_sys}, {"role": "user", "content": prompt_user}])
        return res.choices[0].message.content, None
    except Exception as e: return None, str(e)

# ==============================================================================
# 7. GERADOR PDF (MANTIDO)
# ==============================================================================
class PDF_Classic(FPDF):
    def header(self):
        self.set_draw_color(0, 78, 146); self.set_line_width(0.4)
        self.rect(5, 5, 200, 287)
        logo = finding_logo()
        if logo: self.image(logo, 10, 10, 30); x_offset = 45 
        else: x_offset = 12
        self.set_xy(x_offset, 16); self.set_font('Arial', 'B', 16); self.set_text_color(0, 78, 146)
        self.cell(0, 8, 'PLANO DE ENSINO INDIVIDUALIZADO', 0, 1, 'L')
        self.set_xy(x_offset, 23); self.set_font('Arial', 'I', 10); self.set_text_color(100)
        self.cell(0, 5, 'Documento Oficial de Planejamento Pedagógico', 0, 1, 'L'); self.ln(20)
    def footer(self):
        self.set_y(-15); self.set_font('Arial', 'I', 8); self.set_text_color(128)
        self.cell(0, 10, f'Gerado via PEI 360º | Página {self.page_no()}', 0, 0, 'C')
    def section_title(self, label):
        self.ln(8); self.set_fill_color(240, 248, 255); self.set_text_color(0, 78, 146)
        self.set_font('Arial', 'B', 11); self.cell(0, 8, f"  {label}", 0, 1, 'L', fill=True); self.ln(4)

def gerar_pdf_final(dados, tem_anexo):
    pdf = PDF_Classic(); pdf.add_page(); pdf.set_auto_page_break(auto=True, margin=20)
    pdf.section_title("1. IDENTIFICAÇÃO E CONTEXTO")
    pdf.set_font("Arial", size=10); pdf.set_text_color(0)
    med_list = []
    if dados['lista_medicamentos']:
        for m in dados['lista_medicamentos']:
            obs = m.get('obs', '')
            esc = " (Na Escola)" if m.get('escola') else ""
            txt = f"{m['nome']} ({m['posologia']}){esc}"
            med_list.append(txt)
    med_str = "; ".join(med_list) if med_list else "Não informado."
    fam_str = ", ".join(dados['composicao_familiar_tags']) if dados['composicao_familiar_tags'] else "Não informado."
    pdf.set_font("Arial", 'B', 10); pdf.cell(40, 6, "Nome:", 0, 0); pdf.set_font("Arial", '', 10); pdf.cell(0, 6, dados['nome'], 0, 1)
    pdf.set_font("Arial", 'B', 10); pdf.cell(40, 6, "Nascimento:", 0, 0); pdf.set_font("Arial", '', 10); pdf.cell(0, 6, str(dados['nasc']), 0, 1)
    pdf.set_font("Arial", 'B', 10); pdf.cell(40, 6, "Série/Turma:", 0, 0); pdf.set_font("Arial", '', 10); pdf.cell(0, 6, f"{dados['serie']} - {dados['turma']}", 0, 1)
    pdf.set_font("Arial", 'B', 10); pdf.cell(40, 6, "Diagnóstico:", 0, 0); pdf.set_font("Arial", '', 10); pdf.multi_cell(0, 6, dados['diagnostico']); pdf.ln(2)
    pdf.set_font("Arial", 'B', 10); pdf.cell(40, 6, "Medicação:", 0, 0); pdf.set_font("Arial", '', 10); pdf.multi_cell(0, 6, med_str); pdf.ln(2)
    pdf.set_font("Arial", 'B', 10); pdf.cell(40, 6, "Família:", 0, 0); pdf.set_font("Arial", '', 10); pdf.multi_cell(0, 6, fam_str)
    evid = [k.replace('?', '') for k, v in dados['checklist_evidencias'].items() if v]
    if evid:
        pdf.section_title("2. PONTOS DE ATENÇÃO")
        pdf.set_font("Arial", size=10); pdf.multi_cell(0, 6, limpar_texto_pdf('; '.join(evid) + '.'))
    if any(dados['barreiras_selecionadas'].values()):
        pdf.section_title("3. MAPEAMENTO DE SUPORTE")
        for c, i in dados['barreiras_selecionadas'].items():
            if i:
                pdf.set_font("Arial", 'B', 10); pdf.cell(0, 6, f"{c}:", 0, 1)
                pdf.set_font("Arial", size=10)
                for x in i:
                    niv = dados['niveis_suporte'].get(f"{c}_{x}", "Monitorado")
                    pdf.cell(5); pdf.cell(0, 6, f"- {x}: Suporte {niv}", 0, 1)
                pdf.ln(2)
    if dados['ia_sugestao']:
        pdf.ln(5); pdf.set_text_color(0); pdf.set_font("Arial", '', 10)
        t_limpo = re.sub(r'\[.*?\]', '', dados['ia_sugestao'])
        for linha in t_limpo.split('\n'):
            l = limpar_texto_pdf(linha)
            if re.match(r'^[1-6]\.', l.strip()) and l.strip().isupper():
                pdf.ln(4); pdf.set_fill_color(240, 248, 255); pdf.set_text_color(0, 78, 146); pdf.set_font('Arial', 'B', 11)
                pdf.cell(0, 8, f"  {l}", 0, 1, 'L', fill=True); pdf.set_text_color(0); pdf.set_font("Arial", size=10)
            elif l.strip().endswith(':') and len(l) < 70:
                pdf.ln(2); pdf.set_font("Arial", 'B', 10); pdf.multi_cell(0, 6, l); pdf.set_font("Arial", size=10)
            else: pdf.multi_cell(0, 6, l)
    return pdf.output(dest='S').encode('latin-1', 'replace')

def gerar_docx_final(dados):
    doc = Document(); doc.add_heading('PEI - ' + dados['nome'], 0)
    if dados['ia_sugestao']:
        t_limpo = re.sub(r'\[.*?\]', '', dados['ia_sugestao'])
        doc.add_paragraph(t_limpo)
    b = BytesIO(); doc.save(b); b.seek(0); return b

# ==============================================================================
# 8. INTERFACE UI (PRINCIPAL)
# ==============================================================================
# SIDEBAR
with st.sidebar:
    logo = finding_logo()
    if logo: st.image(logo, width=120)
    if 'OPENAI_API_KEY' in st.secrets: api_key = st.secrets['OPENAI_API_KEY']; st.success("✅ OpenAI OK")
    else: api_key = st.text_input("Chave OpenAI:", type="password")
    st.markdown("### 📂 Carregar Backup")
    uploaded_json = st.file_uploader("Arquivo .json", type="json")
    if uploaded_json:
        try:
            d = json.load(uploaded_json)
            if 'nasc' in d: d['nasc'] = date.fromisoformat(d['nasc'])
            if d.get('monitoramento_data'): d['monitoramento_data'] = date.fromisoformat(d['monitoramento_data'])
            st.session_state.dados.update(d); st.success("Carregado!")
        except: st.error("Erro no arquivo.")
    st.markdown("---")
    if st.button("💾 Salvar no Sistema", use_container_width=True):
        ok, msg = salvar_aluno(st.session_state.dados)
        if ok: st.success(msg)
        else: st.error(msg)
    st.markdown("---")
    data_atual = date.today().strftime("%d/%m/%Y")
    st.markdown(f"<div style='font-size:0.75rem; color:#A0AEC0;'><b>PEI 360º v90.0 Full Text Fix</b><br>Criado e desenvolvido por<br><b>Rodrigo A. Queiroz</b><br>{data_atual}</div>", unsafe_allow_html=True)

# HEADER
logo_path = finding_logo(); b64_logo = get_base64_image(logo_path); mime = "image/png"
img_html = f'<img src="data:{mime};base64,{b64_logo}" style="height: 110px;">' if logo_path else ""

st.markdown(f"""
<div class="header-unified">
    {img_html}
    <div class="header-subtitle">Ecossistema de Inteligência Pedagógica e Inclusiva</div>
</div>""", unsafe_allow_html=True)

# ABAS
abas = ["Início", "Estudante", "Coleta de Evidências", "Rede de Apoio", "Potencialidades & Barreiras", "Plano de Ação", "Monitoramento", "Consultoria IA", "Documento", "🗺️ Meu Mapa"]
tab0, tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab_mapa = st.tabs(abas)

with tab0: # INÍCIO
    if api_key:
        with st.spinner("Gerando inspiração..."):
            try:
                client = OpenAI(api_key=api_key)
                saudacao = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": "Frase curta acolhedora para professor sobre inclusão."}]).choices[0].message.content
                noticia = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": "Dica curta sobre legislação de inclusão ou neurociência."}]).choices[0].message.content
            except:
                saudacao = "A inclusão transforma vidas."
                noticia = "O PEI é um direito garantido por lei."
        st.markdown(f"""
        <div style="background: linear-gradient(90deg, #0F52BA 0%, #004E92 100%); padding: 25px; border-radius: 20px; color: white; margin-bottom: 30px; box-shadow: 0 10px 25px rgba(15, 82, 186, 0.25);">
            <div style="display:flex; gap:20px; align-items:center;">
                <div style="background:rgba(255,255,255,0.2); padding:12px; border-radius:50%;"><i class="ri-sparkling-2-fill" style="font-size: 2rem; color: #FFD700;"></i></div>
                <div><h3 style="color:white; margin:0; font-size: 1.4rem;">Olá, Educador(a)!</h3><p style="margin:5px 0 0 0; opacity:0.95; font-size:1rem;">{saudacao}</p></div>
            </div>
        </div>""", unsafe_allow_html=True)
    
    st.markdown("### <i class='ri-apps-2-line'></i> Fundamentos", unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns(4)
    with c1: st.markdown("""<a href="https://diversa.org.br/educacao-inclusiva/" target="_blank" class="rich-card-link"><div class="home-card hc-blue"><div class="home-icon-box ic-blue"><i class="ri-book-open-line"></i></div><h3>O que é PEI?</h3><p>Conceitos fundamentais da inclusão escolar.</p></div></a>""", unsafe_allow_html=True)
    with c2: st.markdown("""<a href="https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2015/lei/l13146.htm" target="_blank" class="rich-card-link"><div class="home-card hc-gold"><div class="home-icon-box ic-gold"><i class="ri-scales-3-line"></i></div><h3>Legislação</h3><p>Lei Brasileira de Inclusão e Decretos.</p></div></a>""", unsafe_allow_html=True)
    with c3: st.markdown("""<a href="https://institutoneurosaber.com.br/" target="_blank" class="rich-card-link"><div class="home-card hc-pink"><div class="home-icon-box ic-pink"><i class="ri-brain-line"></i></div><h3>Neurociência</h3><p>Artigos sobre desenvolvimento atípico.</p></div></a>""", unsafe_allow_html=True)
    with c4: st.markdown("""<a href="http://basenacionalcomum.mec.gov.br/" target="_blank" class="rich-card-link"><div class="home-card hc-green"><div class="home-icon-box ic-green"><i class="ri-compass-3-line"></i></div><h3>BNCC</h3><p>Currículo oficial e adaptações.</p></div></a>""", unsafe_allow_html=True)
    if api_key: st.markdown(f"""<div class="highlight-card"><i class="ri-lightbulb-flash-fill" style="font-size: 2rem; color: #F59E0B;"></i><div><h4 style="margin:0; color:#1E293B;">Insight de Inclusão</h4><p style="margin:5px 0 0 0; font-size:0.9rem; color:#64748B;">{noticia}</p></div></div>""", unsafe_allow_html=True)

with tab1: # ESTUDANTE
    render_progresso()
    st.markdown("<div class='form-section-title'><i class='ri-user-smile-line'></i> Identidade & Matrícula</div>", unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns([3, 2, 2, 1])
    st.session_state.dados['nome'] = c1.text_input("Nome Completo", st.session_state.dados['nome'])
    st.session_state.dados['nasc'] = c2.date_input("Nascimento", value=st.session_state.dados.get('nasc', date(2015, 1, 1)))
    try: serie_idx = LISTA_SERIES.index(st.session_state.dados['serie']) if st.session_state.dados['serie'] in LISTA_SERIES else 0
    except: serie_idx = 0
    st.session_state.dados['serie'] = c3.selectbox("Série/Ano", LISTA_SERIES, index=serie_idx, placeholder="Selecione...")
    st.session_state.dados['turma'] = c4.text_input("Turma", st.session_state.dados['turma'])
    st.markdown("<div class='form-section-title'><i class='ri-hospital-line'></i> Contexto Clínico & Familiar</div>", unsafe_allow_html=True)
    st.session_state.dados['diagnostico'] = st.text_input("Diagnóstico", st.session_state.dados['diagnostico'])
    c_hist, c_fam = st.columns(2)
    st.session_state.dados['historico'] = c_hist.text_area("Histórico Escolar", st.session_state.dados['historico'])
    st.session_state.dados['familia'] = c_fam.text_area("Dinâmica Familiar", st.session_state.dados['familia'])
    st.session_state.dados['composicao_familiar_tags'] = st.multiselect("Quem mora com o aluno?", LISTA_FAMILIA, default=st.session_state.dados['composicao_familiar_tags'])
    with st.container(border=True):
        usa_med = st.toggle("💊 O aluno faz uso contínuo de medicação?", value=len(st.session_state.dados['lista_medicamentos']) > 0)
        if usa_med:
            c1, c2, c3 = st.columns([3, 2, 2])
            nm = c1.text_input("Nome", key="nm_med")
            pos = c2.text_input("Posologia", key="pos_med")
            admin_escola = c3.checkbox("Administrado na escola?", key="adm_esc")
            if st.button("Adicionar"):
                st.session_state.dados['lista_medicamentos'].append({"nome": nm, "posologia": pos, "obs": "", "escola": admin_escola}); st.rerun()
        if st.session_state.dados['lista_medicamentos']:
            st.write("---")
            for i, m in enumerate(st.session_state.dados['lista_medicamentos']):
                tag = " [NA ESCOLA]" if m.get('escola') else ""
                c_txt, c_btn = st.columns([5, 1])
                c_txt.info(f"💊 **{m['nome']}** ({m['posologia']}){tag}")
                if c_btn.button("Excluir", key=f"del_{i}"): 
                    st.session_state.dados['lista_medicamentos'].pop(i); st.rerun()
    with st.expander("📎 Anexar Laudo (PDF)"):
        up = st.file_uploader("Upload", type="pdf", label_visibility="collapsed")
        if up: st.session_state.pdf_text = ler_pdf(up)

with tab2: # EVIDÊNCIAS
    render_progresso()
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("<div class='form-section-title'><i class='ri-book-open-line'></i> Pedagógico</div>", unsafe_allow_html=True)
        for q in ["Estagnação na aprendizagem", "Dificuldade de generalização", "Dificuldade de abstração", "Lacuna em pré-requisitos"]:
            st.session_state.dados['checklist_evidencias'][q] = st.toggle(q, value=st.session_state.dados['checklist_evidencias'].get(q, False))
    with c2:
        st.markdown("<div class='form-section-title'><i class='ri-brain-line'></i> Cognitivo</div>", unsafe_allow_html=True)
        for q in ["Oscilação de foco", "Fadiga mental rápida", "Dificuldade de iniciar tarefas", "Esquecimento recorrente"]:
            st.session_state.dados['checklist_evidencias'][q] = st.toggle(q, value=st.session_state.dados['checklist_evidencias'].get(q, False))
    with c3:
        st.markdown("<div class='form-section-title'><i class='ri-emotion-line'></i> Comportamental</div>", unsafe_allow_html=True)
        for q in ["Dependência de mediação (1:1)", "Baixa tolerância à frustração", "Desorganização de materiais", "Recusa de tarefas"]:
            st.session_state.dados['checklist_evidencias'][q] = st.toggle(q, value=st.session_state.dados['checklist_evidencias'].get(q, False))

with tab3: # REDE
    render_progresso()
    st.markdown("### <i class='ri-team-line'></i> Rede de Apoio", unsafe_allow_html=True)
    st.session_state.dados['rede_apoio'] = st.multiselect("Profissionais:", LISTA_PROFISSIONAIS, default=st.session_state.dados['rede_apoio'])
    st.session_state.dados['orientacoes_especialistas'] = st.text_area("Orientações Clínicas Importantes", st.session_state.dados['orientacoes_especialistas'])

with tab4: # MAPEAMENTO
    render_progresso()
    with st.container(border=True):
        st.markdown("#### <i class='ri-lightbulb-flash-line' style='color:#0F52BA'></i> Potencialidades e Hiperfoco", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        st.session_state.dados['hiperfoco'] = c1.text_input("Hiperfoco", st.session_state.dados['hiperfoco'])
        p_val = [p for p in st.session_state.dados.get('potencias', []) if p in LISTA_POTENCIAS]
        st.session_state.dados['potencias'] = c2.multiselect("Pontos Fortes", LISTA_POTENCIAS, default=p_val)
    st.divider()
    with st.container(border=True):
        st.markdown("#### <i class='ri-barricade-line' style='color:#FF6B6B'></i> Barreiras e Nível de Suporte", unsafe_allow_html=True)
        c_bar1, c_bar2, c_bar3 = st.columns(3)
        def render_cat_barreira(coluna, titulo, chave_json):
            with coluna:
                st.markdown(f"**{titulo}**")
                itens = LISTAS_BARREIRAS[chave_json]
                b_salvas = [b for b in st.session_state.dados['barreiras_selecionadas'].get(chave_json, []) if b in itens]
                sel = st.multiselect("Selecione:", itens, key=f"ms_{chave_json}", default=b_salvas, label_visibility="collapsed")
                st.session_state.dados['barreiras_selecionadas'][chave_json] = sel
                if sel:
                    for x in sel:
                        st.session_state.dados['niveis_suporte'][f"{chave_json}_{x}"] = st.select_slider(x, ["Autônomo", "Monitorado", "Substancial", "Muito Substancial"], value=st.session_state.dados['niveis_suporte'].get(f"{chave_json}_{x}", "Monitorado"), key=f"sl_{chave_json}_{x}")
                st.write("")
        render_cat_barreira(c_bar1, "Cognitivo", "Cognitivo")
        render_cat_barreira(c_bar1, "Sensorial/Motor", "Sensorial/Motor")
        render_cat_barreira(c_bar2, "Comunicacional", "Comunicacional")
        render_cat_barreira(c_bar2, "Acadêmico", "Acadêmico")
        render_cat_barreira(c_bar3, "Socioemocional", "Socioemocional")

with tab5: # PLANO
    render_progresso()
    st.markdown("### <i class='ri-tools-line'></i> Plano de Ação Estratégico", unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    with c1:
        with st.container(border=True):
            st.markdown("#### 1. Acesso (DUA)")
            st.session_state.dados['estrategias_acesso'] = st.multiselect("Recursos", ["Tempo Estendido", "Apoio Leitura/Escrita", "Material Ampliado", "Tecnologia Assistiva", "Sala Silenciosa"], default=st.session_state.dados['estrategias_acesso'])
            st.session_state.dados['outros_acesso'] = st.text_input("Prática Personalizada (Acesso)", st.session_state.dados['outros_acesso'])
    with c2:
        with st.container(border=True):
            st.markdown("#### 2. Ensino")
            st.session_state.dados['estrategias_ensino'] = st.multiselect("Metodologia", ["Fragmentação de Tarefas", "Pistas Visuais", "Mapas Mentais", "Modelagem", "Ensino Híbrido"], default=st.session_state.dados['estrategias_ensino'])
            st.session_state.dados['outros_ensino'] = st.text_input("Prática Pedagógica (Ensino)", st.session_state.dados['outros_ensino'])
    with c3:
        with st.container(border=True):
            st.markdown("#### 3. Avaliação")
            st.session_state.dados['estrategias_avaliacao'] = st.multiselect("Formato", ["Prova Adaptada", "Prova Oral", "Consulta Permitida", "Portfólio", "Autoavaliação"], default=st.session_state.dados['estrategias_avaliacao'])

with tab6: # MONITORAMENTO
    render_progresso()
    st.markdown("### <i class='ri-loop-right-line'></i> Monitoramento e Metas", unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1: st.session_state.dados['monitoramento_data'] = st.date_input("Próxima Revisão", value=st.session_state.dados.get('monitoramento_data', None))
    with c2: st.session_state.dados['status_meta'] = st.selectbox("Status da Meta Atual", ["Não Iniciado", "Em Andamento", "Parcialmente Atingido", "Atingido", "Superado"], index=0)
    st.write("")
    c3, c4 = st.columns(2)
    with c3: st.session_state.dados['parecer_geral'] = st.selectbox("Parecer Geral", ["Manter Estratégias", "Aumentar Suporte", "Reduzir Suporte (Autonomia)", "Alterar Metodologia", "Encaminhar para Especialista"], index=0)
    with c4: st.session_state.dados['proximos_passos_select'] = st.multiselect("Ações Futuras", ["Reunião com Família", "Encaminhamento Clínico", "Adaptação de Material", "Mudança de Lugar em Sala", "Novo PEI", "Observação em Sala"])

with tab7: # IA (COM NOVO BOTÃO DE REGENERAÇÃO)
    render_progresso()
    st.markdown("### <i class='ri-robot-2-line'></i> Consultoria Pedagógica com IA", unsafe_allow_html=True)
    
    st.markdown("""
    <div style="background-color: #F8FAFC; border-left: 4px solid #0F52BA; padding: 15px; border-radius: 8px; margin-bottom: 20px;">
        <h4 style="color:#0F52BA; margin-top:0;">🤖 Sua Parceira de Planejamento</h4>
        <p style="font-size:0.95rem; color:#4A5568;">
            Olá! Sou sua assistente especialista em <b>Neuroeducação</b> e <b>BNCC</b>. 
            Vou cruzar os dados de <b>Hiperfoco</b>, <b>Barreiras</b> e <b>Laudos</b> 
            para criar um plano de ensino que realmente funciona na prática.
        </p>
    </div>
    """, unsafe_allow_html=True)

    col_left, col_right = st.columns([1, 2])
    with col_left:
        nome_aluno = st.session_state.dados['nome'].split()[0] if st.session_state.dados['nome'] else "o estudante"
        if st.button(f"✨ GERAR PLANO PARA {nome_aluno.upper()}", type="primary", use_container_width=True):
            res, err = consultar_gpt_pedagogico(api_key, st.session_state.dados, st.session_state.pdf_text)
            if res: 
                st.session_state.dados['ia_sugestao'] = res
                effect = random.choice(['balloons', 'snow'])
                if effect == 'balloons': st.balloons()
                else: st.snow()
            else: st.error(err)
        
        # BOTÃO NOVO DE REGENERAÇÃO
        if st.session_state.dados['ia_sugestao']:
            if st.button("🔄 Não gostou? Gerar Nova Abordagem", use_container_width=True):
                 res, err = consultar_gpt_pedagogico(api_key, st.session_state.dados, st.session_state.pdf_text, regenerar=True)
                 if res: st.session_state.dados['ia_sugestao'] = res; st.rerun()

        with st.expander("📚 Base Técnica & Legal"):
            st.markdown("""
            **1. Documentos Norteadores**
            * NOTA TÉCNICA SEESP/MEC nº 24/2010.
            * DUA - Desenho Universal para a Aprendizagem.
            
            **2. Autores de Referência**
            * MENDES, Enicéia Gonçalves (Ensino Colaborativo).
            * MANTOAN, Maria Teresa Eglér (Inclusão Total).
            """)

    with col_right:
        if st.session_state.dados['ia_sugestao']:
            with st.expander("🔍 Entenda a Lógica (Calibragem)"):
                st.markdown("""**Como este plano foi construído:**\n* **Filtro Vygotsky:** Identificação da Zona de Desenvolvimento Proximal.\n* **Análise Farmacológica:** Impacto da medicação na aprendizagem.""")
            st.markdown(st.session_state.dados['ia_sugestao'])
            st.info("📝 **Personalize:** O texto acima é editável.")
            novo_texto = st.text_area("Editor de Conteúdo", value=st.session_state.dados['ia_sugestao'], height=400, key="editor_ia")
            st.session_state.dados['ia_sugestao'] = novo_texto
        else:
            st.info(f"👈 Clique no botão ao lado para gerar o plano de {nome_aluno}.")

with tab8: # DASHBOARD FINAL (V74)
    render_progresso()
    st.markdown("### <i class='ri-file-pdf-line'></i> Dashboard e Exportação", unsafe_allow_html=True)
    if st.session_state.dados['nome']:
        init_avatar = st.session_state.dados['nome'][0].upper() if st.session_state.dados['nome'] else "?"
        idade_str = calcular_idade(st.session_state.dados['nasc'])
        
        st.markdown(f"""
        <div class="dash-hero">
            <div style="display:flex; align-items:center; gap:20px;">
                <div class="apple-avatar">{init_avatar}</div>
                <div style="color:white;"><h1>{st.session_state.dados['nome']}</h1><p>{st.session_state.dados['serie']}</p></div>
            </div>
            <div>
                <div style="text-align:right; font-size:0.8rem; opacity:0.8;">IDADE</div>
                <div style="font-size:1.2rem; font-weight:bold;">{idade_str}</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        c_kpi1, c_kpi2, c_kpi3, c_kpi4 = st.columns(4)
        with c_kpi1:
            n_pot = len(st.session_state.dados['potencias'])
            color_p = "#38A169" if n_pot > 0 else "#CBD5E0"
            st.markdown(f"""<div class="metric-card"><div class="css-donut" style="--p: {n_pot*10}%; --fill: {color_p};"><div class="d-val">{n_pot}</div></div><div class="d-lbl">Potencialidades</div></div>""", unsafe_allow_html=True)
        with c_kpi2:
            n_bar = sum(len(v) for v in st.session_state.dados['barreiras_selecionadas'].values())
            color_b = "#E53E3E" if n_bar > 5 else "#DD6B20"
            st.markdown(f"""<div class="metric-card"><div class="css-donut" style="--p: {n_bar*5}%; --fill: {color_b};"><div class="d-val">{n_bar}</div></div><div class="d-lbl">Barreiras</div></div>""", unsafe_allow_html=True)
        with c_kpi3:
             hf = st.session_state.dados['hiperfoco'] or "-"
             hf_emoji = get_hiperfoco_emoji(hf)
             st.markdown(f"""<div class="metric-card"><div style="font-size:2.5rem;">{hf_emoji}</div><div style="font-weight:800; font-size:1.1rem; color:#2D3748; margin:10px 0;">{hf}</div><div class="d-lbl">Hiperfoco</div></div>""", unsafe_allow_html=True)
        with c_kpi4:
             txt_comp, bg_c, txt_c = calcular_complexidade_pei(st.session_state.dados)
             st.markdown(f"""<div class="metric-card" style="background-color:{bg_c}; border-color:{txt_c};"><div class="comp-icon-box"><i class="ri-error-warning-line" style="color:{txt_c}; font-size: 2rem;"></i></div><div style="font-weight:800; font-size:1.1rem; color:{txt_c}; margin:5px 0;">{txt_comp}</div><div class="d-lbl" style="color:{txt_c};">Nível de Atenção</div></div>""", unsafe_allow_html=True)

        st.write("")
        c_r1, c_r2 = st.columns(2)
        with c_r1:
            tem_med = len(st.session_state.dados['lista_medicamentos']) > 0
            if tem_med:
                st.markdown(f"""<div class="soft-card sc-orange"><div class="sc-head"><i class="ri-medicine-bottle-fill" style="color:#DD6B20;"></i> Atenção Farmacológica</div><div class="sc-body">Aluno em uso de medicação contínua. Verifique a aba Estudante para detalhes e posologia.</div><div class="bg-icon">💊</div></div>""", unsafe_allow_html=True)
            else:
                st.markdown(f"""<div class="soft-card sc-green"><div class="sc-head"><i class="ri-checkbox-circle-fill" style="color:#38A169;"></i> Medicação</div><div class="sc-body">Nenhuma medicação informada.</div><div class="bg-icon">✅</div></div>""", unsafe_allow_html=True)
            st.write("")
            metas = extrair_metas_estruturadas(st.session_state.dados['ia_sugestao'])
            if metas:
                html_metas = f"""<div class="meta-row"><span style="font-size:1.2rem;">🏁</span> <b>Curto (2m):</b> {metas['Curto']}</div><div class="meta-row"><span style="font-size:1.2rem;">🧗</span> <b>Médio (6m):</b> {metas['Medio']}</div><div class="meta-row"><span style="font-size:1.2rem;">🏔️</span> <b>Longo (1a):</b> {metas['Longo']}</div>"""
            else: html_metas = "Gere o plano na aba IA."
            st.markdown(f"""<div class="soft-card sc-yellow"><div class="sc-head"><i class="ri-flag-2-fill" style="color:#D69E2E;"></i> Cronograma de Metas</div><div class="sc-body">{html_metas}</div></div>""", unsafe_allow_html=True)

        with c_r2:
            verbos = extrair_bloom(st.session_state.dados['ia_sugestao'])
            html_verbos = "".join([f'<span class="bloom-tag">{v}</span>' for v in verbos])
            st.markdown(f"""<div class="soft-card sc-blue"><div class="sc-head"><i class="ri-lightbulb-flash-fill" style="color:#3182CE;"></i> Taxonomia de Bloom (Verbos)</div><div class="sc-body"><div style="margin-bottom:10px; font-size:0.85rem; color:#4A5568;">Verbos de comando sugeridos para atividades:</div>{html_verbos}</div><div class="bg-icon">🧠</div></div>""", unsafe_allow_html=True)
            st.write("")
            rede_html = ""
            if st.session_state.dados['rede_apoio']:
                for prof in st.session_state.dados['rede_apoio']:
                    icon = get_pro_icon(prof)
                    rede_html += f'<span class="rede-chip">{icon} {prof}</span> '
            else: rede_html = "<span style='opacity:0.6;'>Sem rede de apoio.</span>"
            st.markdown(f"""<div class="soft-card sc-cyan"><div class="sc-head"><i class="ri-team-fill" style="color:#0BC5EA;"></i> Rede de Apoio</div><div class="sc-body">{rede_html}</div><div class="bg-icon">🤝</div></div>""", unsafe_allow_html=True)

        st.write("")
        st.markdown("##### 🧬 DNA de Suporte (Detalhamento)")
        st.markdown('<div class="dna-legend"><i class="ri-information-fill"></i> Barras maiores indicam áreas que exigem mais adaptação e suporte intenso.</div>', unsafe_allow_html=True)
        dna_c1, dna_c2 = st.columns(2)
        areas = list(LISTAS_BARREIRAS.keys())
        for i, area in enumerate(areas):
            qtd = len(st.session_state.dados['barreiras_selecionadas'].get(area, []))
            val = min(qtd * 20, 100)
            target = dna_c1 if i < 3 else dna_c2
            color = "#3182CE"
            if val > 40: color = "#DD6B20"
            if val > 70: color = "#E53E3E"
            target.markdown(f"""<div class="dna-bar-container"><div class="dna-bar-flex"><span>{area}</span><span>{qtd} barreiras</span></div><div class="dna-bar-bg"><div class="dna-bar-fill" style="width:{val}%; background:{color};"></div></div></div>""", unsafe_allow_html=True)

    st.divider()
    if st.session_state.dados['ia_sugestao']:
        c1, c2 = st.columns(2)
        with c1:
            pdf = gerar_pdf_final(st.session_state.dados, len(st.session_state.pdf_text)>0)
            st.download_button("📥 Baixar PDF Oficial", pdf, f"PEI_{st.session_state.dados['nome']}.pdf", "application/pdf", type="primary")
        with c2:
            docx = gerar_docx_final(st.session_state.dados)
            st.download_button("📥 Baixar Word Editável", docx, f"PEI_{st.session_state.dados['nome']}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.write("")
            json_dados = json.dumps(st.session_state.dados, default=str)
            st.download_button("💾 Baixar Arquivo do Aluno (.json)", json_dados, f"PEI_{st.session_state.dados['nome']}.json", "application/json")

with tab_mapa: # MAPA FINAL
    render_progresso()
    st.markdown(f"""
    <div style="background: linear-gradient(90deg, #F6E05E 0%, #D69E2E 100%); padding: 25px; border-radius: 20px; color: #2D3748; margin-bottom: 20px; box-shadow: 0 4px 6px rgba(0,0,0,0.05);">
        <h3 style="margin:0; color:#2D3748;">🗺️ Meu Mapa da Jornada</h3>
        <p style="margin:5px 0 0 0; font-weight:600;">Olá, {st.session_state.dados['nome'].split()[0] if st.session_state.dados['nome'] else 'Estudante'}! Este é o seu plano de jogo para aprender melhor.</p>
    </div>
    """, unsafe_allow_html=True)
    
    col_text_map, col_dalle_map = st.columns([1.5, 2])
    
    # TEXTO SIMPLES E CLARO
    with col_text_map:
        st.markdown("#### ⚡ Meus Poderes & Missões")
        if st.session_state.dados['ia_sugestao']:
            texto_mapa = extrair_tag_ia(st.session_state.dados['ia_sugestao'], "MAPA_TEXTO_GAMIFICADO")
            if texto_mapa:
                with st.container(border=True):
                    st.markdown(texto_mapa)
            else:
                st.warning("O mapa de texto ainda não foi gerado. Clique em 'Gerar Plano' na aba IA.")
        else:
            st.info("Gere o plano na aba IA primeiro.")

    # IMAGEM VISUAL (GERADA PELO TEXTO)
    with col_dalle_map:
        st.markdown("#### 🎨 Meu Quadro Visual (DALL-E)")
        st.markdown("""<p style="font-size:0.85rem; color:#718096;">Gera um infográfico estilo 'Mapa Mental' baseado exatamente no texto ao lado.</p>""", unsafe_allow_html=True)
        
        texto_para_imagem = ""
        if st.session_state.dados['ia_sugestao']:
             texto_para_imagem = extrair_tag_ia(st.session_state.dados['ia_sugestao'], "MAPA_TEXTO_GAMIFICADO")

        if st.button("✨ Criar Mapa Visual (Baseado no Texto)", type="primary", use_container_width=True):
            if texto_para_imagem and st.session_state.dados['hiperfoco']:
                # Chama a função integrada (prompt com texto)
                url, err = gerar_imagem_dalle_integrada(api_key, st.session_state.dados, texto_para_imagem)
                if url:
                    st.session_state.dalle_image_url = url
                    st.success("Mapa visual gerado com sucesso!")
                else:
                    st.error(f"Erro ao gerar imagem: {err}")
            else:
                st.warning("Certifique-se de que o plano de texto foi gerado e o Hiperfoco está definido.")
        
        if st.session_state.dalle_image_url:
            st.image(st.session_state.dalle_image_url, use_column_width=True, caption="Visualização das suas estratégias")
            st.markdown(f'<a href="{st.session_state.dalle_image_url}" download="Mapa_Visual_Integrado.png" target="_blank" style="display:block; text-decoration:none; background-color:#0F52BA; color:white; padding:10px; border-radius:8px; text-align:center; margin-top:10px;">📥 Baixar Imagem do Mapa</a>', unsafe_allow_html=True)

st.markdown("---")
